#!/usr/bin/env python3
'''creatiion of sticker in microsoft document'''

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm
from load_data import load_data

def sticker_doc_creation():
    '''Method to create sticker in docx'''
    items, weights, mrps = load_data()
    for item in items:
    # Create a new Document
    doc = Document()

    # Add a table
    table = doc.add_table(rows=8, cols=5)

    # Set table width
    table.autofit = True  # Autofit the table to the page width

    # Set minimum row height
    for row in table.rows:
        row.height = Inches(0.4)  # Adjust the height as needed (e.g., Inches(0.5) for 0.5 inches)

    # Set minimum column width
    table.allow_autofit = False  # Disable autofitting to set minimum column width
    for column in table.columns:
        column.width = Inches(1.6)  # Adjust the width as needed (e.g., Inches(1.5) for 1.5 inches)

    # Set page margins in inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.4)      # 0.4 centimeters
        section.bottom_margin = Cm(0)      # 0 centimeters (no margin) at the bottom
        section.left_margin = Cm(0.4)      # 0.4 centimeters
        section.right_margin = Cm(0.4)     # 0.4 centimeters

    data = {
        'item': 'ಜೀರಿಗೆ',
        'mrp': 'Rs.75/-',
        'net_weight': '250 gms',
        'pkg_date': 'Feb.2017',
        'fssai_license': '11216336000104',
        'batch_no': 'AF-02/17-18'
    }

    # Sample multi-line text with placeholders for parameterization
    multi_line_text = f"""
    AMRUTA FOODS 
    Item: {data['item']}
    MRP: {data['mrp']}
    Incl. of all taxes
    Net wt: {data['net_weight']}
    Pkd. date: {data['pkg_date']}
    fssai lic.no. {data['fssai_license']}
    Batch no. {data['batch_no']}"""

    # Iterate over each row in the table
    for row in table.rows:
        # Iterate over each cell in the row
        for cell in row.cells:
            # Create a new paragraph in the cell
            pt = cell.paragraphs[0]
            # t = pt.text = ''
            pt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Split the multi-line text into lines
            lines = multi_line_text.strip().split('\n')
            count = 0

            # Add each line to the paragraph
            for line in lines:
                # Create a new run for the line
                # run = paragraph.add_run(line)
                run = pt.add_run(line)

                # Set font properties
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8.5)
                run.space_before = Pt(0)
                run.space_after = Pt(0)

                # Set bold font for the header line
                if 'AMRUTA FOODS' in line:
                    run.bold = True
                    run.underline = True
                    run.font.size = Pt(9.5)
                    run.space_before = Pt(0)
                    run.space_after = Pt(0)

                if 'fssai lic.no.' in line:
                    # run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.italic = True
                    run.space_before = Pt(0)
                    run.space_after = Pt(0)

                if count !=  (len(lines)-1):
                    run = pt.add_run('\n')
                count += 1

    # Save the document
    doc.save('item.docx')
