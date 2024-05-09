#!/usr/bin/env python3
'''creatiion of sticker in microsoft document'''

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm
from load_data import load_data
import utils

def sticker_doc_creation():
    '''Method to create sticker in docx'''
    items_english, items, weights, mrps = load_data()
    month, year = utils.get_current_date()
    batch_no = utils.batch_number()
    directory = 'Amruta Foods'
    utils.folder_creation(directory)
    os.chdir(directory)
    for item_english, item, weight_list, mrp_list in zip(items_english, items, weights, mrps):
        os.mkdir(item_english)
        os.chdir(item_english)
        for weight, mrp in zip(weight_list, mrp_list):
            # Create a new Document
            doc = Document()
            doc_name = weight + ".docx"

            if weight == '1':
                weight = weight + " kg"
            else:
                weight = weight + " gms"

            # Add a table
            table = doc.add_table(rows=8, cols=5)

            # Set table width
            table.autofit = True  # Autofit the table to the page width

            # Set minimum row height
            for row in table.rows:
                row.height = Inches(0.4)  # Adjust the height as needed

            # Set minimum column width
            table.allow_autofit = False  # Disable autofitting to set minimum column width
            for column in table.columns:
                column.width = Inches(1.6)  # Adjust the width as needed

            # Set page margins in inches
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(0.3)
                section.bottom_margin = Cm(0)
                section.left_margin = Cm(0.45)
                section.right_margin = Cm(0.45)

            cell_data = f"""
AMRUTA FOODS 
Item: {item}
MRP: Rs.{mrp}/-
Incl. of all taxes
Net wt: {weight}
Pkd. date: {month}.{year}
fssai lic.no. 11216336000104
Batch no. {batch_no}"""

            # Iterate over each row in the table
            for row in table.rows:
                # Iterate over each cell in the row
                for cell in row.cells:
                    # Create a new paragraph in the cell
                    pt = cell.paragraphs[0]
                    pt.paragraph_format.space_after = Pt(5)
                    pt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Split the multi-line text into lines
                    lines = cell_data.strip().split('\n')
                    count = 0

                    # Add each line to the paragraph
                    for line in lines:
                        # Create a new run for the line
                        # run = paragraph.add_run(line)
                        run = pt.add_run(line)

                        # Set font properties
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8.5)

                        # Set bold font for the header line
                        if 'AMRUTA FOODS' in line:
                            run.bold = True
                            run.underline = True
                            run.font.size = Pt(9.5)

                        if 'fssai lic.no.' in line:
                            # run.font.name = 'Arial'
                            run.font.size = Pt(8)
                            run.italic = True

                        if count !=  (len(lines)-1):
                            run = pt.add_run('\n')
                        count += 1

            # Save the document
            doc.save(doc_name)
        os.chdir("..")
